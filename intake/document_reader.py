from __future__ import annotations

from pathlib import Path


class DocumentReadError(RuntimeError):
    """Raised when an input document cannot be read."""


def read_document(path: str | Path) -> dict:
    """Read supported intake documents into plain text.

    Supported formats are docx, pdf, txt, and md. The docx/pdf dependencies are
    optional so the basic text workflow remains lightweight.
    """

    source = Path(path)
    suffix = source.suffix.lower()
    if not source.exists():
        raise DocumentReadError(f"Input file does not exist: {source}")

    format_clues = empty_format_clues()
    warnings: list[str] = []
    if suffix in {".txt", ".md"}:
        text = _read_text(source)
    elif suffix == ".docx":
        text, format_clues = _read_docx(source)
    elif suffix == ".pdf":
        text = _read_pdf(source)
        warnings.append(
            "PDF extraction does not reliably preserve bold, underline, or starred author formatting. "
            "Provide a Word file or manually confirm target/corresponding authors."
        )
    else:
        raise DocumentReadError(f"Unsupported input format: {suffix or '<none>'}")

    return {
        "path": str(source),
        "format": suffix.lstrip("."),
        "text": text,
        "line_count": len([line for line in text.splitlines() if line.strip()]),
        "format_clues": format_clues,
        "warnings": warnings,
    }


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError(f"Could not decode text file: {path}")


def _read_docx(path: Path) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentReadError("Reading .docx files requires python-docx.") from exc

    document = Document(str(path))
    parts: list[str] = []
    format_clues = empty_format_clues()
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text:
            parts.append(paragraph.text)
            collect_run_formatting(paragraph.runs, paragraph_text, format_clues)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_run_formatting(paragraph.runs, paragraph.text, format_clues)
    return "\n".join(parts), dedupe_format_clues(format_clues)


def _read_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise DocumentReadError("Reading .pdf files requires pdfplumber.") from exc

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
    return "\n\n".join(parts)


def empty_format_clues() -> dict:
    return {
        "bold_segments": [],
        "underlined_segments": [],
        "starred_segments": [],
    }


def collect_run_formatting(runs, paragraph_text: str, format_clues: dict) -> None:
    for run in runs:
        text = clean_segment(run.text)
        if not looks_like_author_segment(text):
            continue
        if run.bold:
            format_clues["bold_segments"].append(text)
        if run.underline:
            format_clues["underlined_segments"].append(text)
    for segment in extract_starred_segments(paragraph_text):
        format_clues["starred_segments"].append(segment)


def extract_starred_segments(text: str) -> list[str]:
    import re

    results = []
    for match in re.finditer(r"([A-Z][A-Za-z.\-'\s]{1,50})\*", text or ""):
        segment = clean_segment(match.group(1))
        if looks_like_author_segment(segment):
            results.append(segment)
    return results


def clean_segment(value: str) -> str:
    return " ".join((value or "").replace("\n", " ").split()).strip(" ,;*")


def looks_like_author_segment(value: str) -> bool:
    if not value or len(value) > 80:
        return False
    if any(char.isdigit() for char in value):
        return False
    if len(value.split()) > 6:
        return False
    return any(char.isalpha() for char in value)


def dedupe_format_clues(format_clues: dict) -> dict:
    return {
        key: sorted(set(value for value in values if value))
        for key, values in format_clues.items()
    }
