import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from intake.document_reader import read_document
from intake.profile_extractor import extract_profile


class TestDocumentIntake(unittest.TestCase):
    def test_text_profile_extraction(self):
        text = """Name: Xiaolu Li
Affiliation: IMDEA Materials Institute
Email: xiaolu.li@example.edu
ORCID: 0000-0002-1825-0097
Google Scholar: https://scholar.google.com/citations?user=abcDEF1AAAAJ&hl=en

Publications
1. Xiaolu Li, A. Smith. Fire behavior of polymer composites. Advanced Materials, 2024. DOI: 10.1002/adma.202400001
"""
        profile = extract_profile(text)
        self.assertEqual(profile["name"], "Xiaolu Li")
        self.assertEqual(profile["affiliation"], "IMDEA Materials Institute")
        self.assertEqual(profile["orcid"], "0000-0002-1825-0097")
        self.assertEqual(profile["scholar_user_id"], "abcDEF1AAAAJ")
        self.assertIn("10.1002/adma.202400001", profile["dois"])
        self.assertEqual(profile["publications"][0]["year"], "2024")

    def test_read_txt_and_md(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = Path(tmp) / "cv.txt"
            md = Path(tmp) / "cv.md"
            txt.write_text("Name: Test Scholar", encoding="utf-8")
            md.write_text("# CV\nName: Test Scholar", encoding="utf-8")
            self.assertEqual(read_document(txt)["format"], "txt")
            self.assertEqual(read_document(md)["format"], "md")

    def test_read_docx_when_dependency_available(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.docx"
            doc = Document()
            doc.add_paragraph("Name: Test Scholar")
            doc.save(path)
            result = read_document(path)
            self.assertIn("Test Scholar", result["text"])

    def test_read_docx_author_formatting_when_dependency_available(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            bold = paragraph.add_run("Xiaolu Li")
            bold.bold = True
            paragraph.add_run(", ")
            underlined = paragraph.add_run("A. Smith")
            underlined.underline = True
            paragraph.add_run("* Fire behavior of polymer composites. Advanced Materials, 2024.")
            doc.save(path)
            result = read_document(path)
            self.assertIn("Xiaolu Li", result["format_clues"]["bold_segments"])
            self.assertIn("A. Smith", result["format_clues"]["underlined_segments"])
            self.assertIn("A. Smith", result["format_clues"]["starred_segments"])

    def test_read_pdf_when_dependency_available(self):
        try:
            import pdfplumber  # noqa: F401
            from reportlab.pdfgen import canvas
        except ImportError:
            self.skipTest("pdfplumber/reportlab is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.pdf"
            pdf = canvas.Canvas(str(path))
            pdf.drawString(72, 720, "Name: Test Scholar")
            pdf.save()
            result = read_document(path)
            self.assertIn("Test Scholar", result["text"])
            self.assertTrue(result["warnings"])


if __name__ == "__main__":
    unittest.main()
